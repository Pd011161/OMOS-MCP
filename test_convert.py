"""Self-check for the file converters and thread safety. Run: uv run python test_convert.py"""
import io
import threading
import time
from unittest.mock import patch

from omos_mcp import server
from omos_mcp.server import _read_docx, _read_xlsx, _truncate

# docx
import docx

d = docx.Document()
d.add_paragraph("hello world")
t = d.add_table(rows=1, cols=2)
t.rows[0].cells[0].text = "a"
t.rows[0].cells[1].text = "b"
buf = io.BytesIO()
d.save(buf)
out = _read_docx(buf.getvalue())
assert "hello world" in out and "a | b" in out, out

# xlsx
import openpyxl

wb = openpyxl.Workbook()
wb.active.title = "S1"
wb.active.append(["x", 1, None])
buf = io.BytesIO()
wb.save(buf)
out = _read_xlsx(buf.getvalue())
assert "Sheet: S1" in out and "x | 1 |" in out, out

# truncate
assert _truncate("") == "(file is empty)"
assert _truncate("a" * 60_000).endswith("chars total]")

# thread safety: one Drive client per thread (sharing one breaks TLS), index built once
builds = []


def _fake_build(*args, **kwargs):
    client = object()
    builds.append(client)
    return client


with patch.object(server, "OMOS_ROOT_FOLDER_ID", "x"), \
     patch.object(server, "GOOGLE_SERVICE_ACCOUNT_JSON", "{}"), \
     patch("google.oauth2.service_account.Credentials.from_service_account_info", lambda *a, **k: None), \
     patch("googleapiclient.discovery.build", _fake_build):
    def grab():
        server._svc()
        server._svc()  # must reuse this thread's client

    threads = [threading.Thread(target=grab) for _ in range(4)]
    [t.start() for t in threads]
    [t.join() for t in threads]
assert len(builds) == 4, f"expected 1 client per thread (4), got {len(builds)}"

# concurrent callers must not each rebuild the project list
calls = []


def _slow_list(folder_id):
    calls.append(folder_id)
    time.sleep(0.1)
    return []


server._projects.update(built_at=0.0, by_name={})
with patch.object(server, "_list_children", _slow_list):
    threads = [threading.Thread(target=server._ensure_projects) for _ in range(4)]
    [t.start() for t in threads]
    [t.join() for t in threads]
assert len(calls) == 1, f"project list built {len(calls)} times, expected 1"

# project list: one API call, names only — the drive has hundreds of projects and
# tens of thousands of files, so nothing may walk it whole
ROOT = "root"
_TREE = {
    ROOT: [
        {"id": "p1", "name": "Alpha", "mimeType": server.FOLDER_MT, "webViewLink": "l1", "parents": [ROOT]},
        {"id": "p2", "name": "Beta", "mimeType": server.FOLDER_MT, "webViewLink": "l2", "parents": [ROOT]},
        {"id": "f0", "name": "loose.md", "mimeType": "text/markdown", "webViewLink": "l0", "parents": [ROOT]},
    ],
    "p1": [
        {"id": "d1", "name": "Docs", "mimeType": server.FOLDER_MT, "webViewLink": "l3", "parents": ["p1"]},
        {"id": "f1", "name": "top.md", "mimeType": "text/markdown", "webViewLink": "l4", "parents": ["p1"]},
    ],
    "d1": [{"id": "f2", "name": "deep.pdf", "mimeType": "application/pdf", "webViewLink": "l5", "parents": ["d1"]}],
    "p2": [],
}
listed = []


def _fake_children(folder_id):
    listed.append(folder_id)
    return _TREE.get(folder_id, [])


server._projects.update(built_at=0.0, by_name={})
server._folders.clear()
with patch.object(server, "OMOS_ROOT_FOLDER_ID", ROOT), \
     patch.object(server, "_list_children", _fake_children):
    out = server._omos_index()
    assert listed == [ROOT], f"index must hit the drive once, hit: {listed}"
    assert "Alpha" in out and "Beta" in out, out
    assert "deep.pdf" not in out and "loose.md" not in out, "index must list project names only"

    # one project's subtree only — Beta is never touched
    listed.clear()
    out = server._omos_list("Alpha", "")
    assert "top.md" in out and "deep.pdf" in out and "📁 Docs" in out, out
    assert "p2" not in listed, f"listing Alpha must not walk other projects: {listed}"

    out = server._omos_list("Nope", "")
    assert "Unknown project" in out, out
    assert "deep.pdf" in server._omos_list("Alpha", "Docs"), "subfolder listing failed"

    # search resolves each hit's path by walking parents, no full index
    with patch.object(server, "_svc", lambda: None):
        path, project = server._path_of("deep.pdf", ["d1"])
    assert (path, project) == ("Alpha / Docs / deep.pdf", "Alpha"), (path, project)
    assert server._cite("deep.pdf", ["d1"], "l5").startswith("📄 **deep.pdf**")

# ponytail: pdf extraction not self-checked — pypdf can't author text PDFs; covered by real-drive test
print("ok")
