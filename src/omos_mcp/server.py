import io
import json
import logging
import os
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from mcp.server.fastmcp import FastMCP, Image


def _load_dotenv() -> None:
    # ponytail: stdlib .env loader (repo root or cwd), real env always wins
    for p in (Path(__file__).resolve().parents[2] / ".env", Path(".env")):
        if p.is_file():
            for line in p.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip().strip("'\""))
            return


_load_dotenv()

# --- Configuration (all overridable via environment) ---
OMOS_ROOT_FOLDER_ID = os.environ.get("OMOS_ROOT_FOLDER_ID", "")
GOOGLE_SERVICE_ACCOUNT_JSON = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "")
INDEX_TTL = int(os.environ.get("OMOS_INDEX_TTL", "300"))  # seconds
# A tool call must always come back with something: clients cut the call off long
# before a big drive walk finishes, and an error is worth less than partial results.
DEADLINE = float(os.environ.get("OMOS_DEADLINE", "20"))  # seconds per tool call
HTTP_TIMEOUT = float(os.environ.get("OMOS_HTTP_TIMEOUT", "20"))  # seconds per Drive request

FOLDER_MT = "application/vnd.google-apps.folder"
MAX_DOWNLOAD_BYTES = 20 * 1024 * 1024
MAX_TEXT_CHARS = 50_000

INSTRUCTIONS = (
    "OMOS shared-drive knowledge base — the single source of truth for ALL internal project "
    "documents. USE THESE TOOLS (without being asked) whenever the user asks about any project, "
    "BRD, requirement, timeline, แผนงาน, กำหนดการ, system flow, database/DB schema, API spec, "
    "project overview, รายละเอียดโปรเจค, เอกสารโปรเจค, or mentions a project by name.\n"
    "Each top-level folder is one project — the folder name IS the project name. The layout "
    "inside a project varies; there is no fixed structure to assume.\n\n"
    "The drive holds hundreds of projects and tens of thousands of files, so work top-down:\n"
    "1. omos_index(filter=<keyword the user said>) — finds projects by name. ALWAYS start here, "
    "and ALWAYS pass a filter when the user named or hinted at a project: without one you get a "
    "small sample, not the whole drive, so never conclude a project is missing from that sample. "
    "Try a shorter keyword before giving up.\n"
    "2. Match the user's wording against the names you get back. If nothing matches clearly, pick "
    "the likeliest candidates and ASK THE USER TO CONFIRM; if you cannot guess, ask for the "
    "project name. Never guess silently.\n"
    "3. omos_list(project) — the files inside that project (add subfolder= to narrow a big one).\n"
    "4. omos_search(query, project) — keyword search, scoped to a project whenever you know it.\n"
    "5. omos_read(file_id) — read a file you picked in step 3 or 4.\n"
    "If the question is too broad, ask the user to narrow it before searching. "
    "EVERY answer must cite the source file(s) by name WITH the Google Drive link "
    "(provided in every tool response)."
)

mcp = FastMCP("omos-mcp", instructions=INSTRUCTIONS)


# --- Google Drive client ---

# google-api-python-client is not thread-safe: sharing one client across the
# worker threads corrupts the TLS connection ("[SSL] record layer failure"),
# so each thread gets its own.
_local = threading.local()


def _svc():
    drive = getattr(_local, "drive", None)
    if drive is None:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build

        if not OMOS_ROOT_FOLDER_ID:
            raise RuntimeError("OMOS_ROOT_FOLDER_ID is not set.")
        raw = GOOGLE_SERVICE_ACCOUNT_JSON
        if not raw:
            raise RuntimeError("GOOGLE_SERVICE_ACCOUNT_JSON is not set (JSON content or a file path).")
        info = json.loads(raw) if raw.lstrip().startswith("{") else json.loads(
            Path(raw).expanduser().read_text(encoding="utf-8")
        )
        creds = service_account.Credentials.from_service_account_info(
            info, scopes=["https://www.googleapis.com/auth/drive.readonly"]
        )
        import google_auth_httplib2
        import httplib2

        drive = build(
            "drive", "v3", cache_discovery=False,
            http=google_auth_httplib2.AuthorizedHttp(
                creds, http=httplib2.Http(timeout=HTTP_TIMEOUT)
            ),
        )
        _local.drive = drive
    return drive


_FILE_FIELDS = "id,name,mimeType,webViewLink,parents"


def _list_children(folder_id: str) -> list[dict]:
    files, token = [], None
    while True:
        resp = _svc().files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields=f"nextPageToken, files({_FILE_FIELDS})",
            pageSize=1000,
            pageToken=token,
            orderBy="folder,name",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        files += resp.get("files", [])
        token = resp.get("nextPageToken")
        if not token:
            return files


# --- Caches ---
# The drive holds hundreds of projects and tens of thousands of files, so nothing
# walks it whole: the project list is one API call, and paths are resolved only
# for the handful of files a search or read actually touches.
_projects: dict = {"built_at": 0.0, "by_name": {}}  # name -> [folder ids]
_folders: dict[str, tuple[str, str]] = {}  # folder id -> (name, parent id)
_lock = threading.Lock()


def _ensure_projects() -> None:
    if time.time() - _projects["built_at"] <= INDEX_TTL:
        return
    with _lock:
        if time.time() - _projects["built_at"] <= INDEX_TTL:
            return  # another thread just built it
        started = time.time()
        by_name: dict[str, list[str]] = {}
        for f in _list_children(OMOS_ROOT_FOLDER_ID):
            if f["mimeType"] == FOLDER_MT:
                by_name.setdefault(f["name"].strip(), []).append(f["id"])
                _folders[f["id"]] = (f["name"].strip(), OMOS_ROOT_FOLDER_ID)
        _projects.update(built_at=time.time(), by_name=by_name)
        logging.getLogger("omos_mcp").info(
            "project list built in %.1fs: %d project(s)", time.time() - started, len(by_name)
        )


def _folder(folder_id: str) -> tuple[str, str]:
    """(name, parent id) for a folder, cached — folder names never change often."""
    hit = _folders.get(folder_id)
    if hit is None:
        meta = _svc().files().get(
            fileId=folder_id, fields="id,name,parents", supportsAllDrives=True
        ).execute()
        parents = meta.get("parents") or [""]
        hit = (meta.get("name", "?"), parents[0])
        _folders[folder_id] = hit
    return hit


def _path_of(name: str, parents: list[str] | None) -> tuple[str, str]:
    """Resolve ("Project / Sub / file.md", "Project") by walking parents up to the root."""
    chain: list[str] = []
    fid = (parents or [""])[0]
    for _ in range(8):  # ponytail: depth cap, matches the drive's real nesting
        if not fid or fid == OMOS_ROOT_FOLDER_ID:
            break
        folder_name, fid = _folder(fid)
        chain.append(folder_name)
    chain.reverse()
    project = chain[0] if chain else ""
    return " / ".join([*chain, name]), project


def _cite(name: str, parents: list[str] | None, link: str) -> str:
    path, _ = _path_of(name, parents)
    return f"📄 **{name}**\n📁 {path}\n🔗 {link}\n\n---\n\n"


# --- Tools ---
# All tools are async and push the blocking Google API work into a worker thread:
# FastMCP runs sync tools directly on the event loop, which froze /healthz during
# long Drive walks and made Render kill the instance mid-call.


@mcp.tool()
async def omos_index(filter: str = "", limit: int = 10) -> str:
    """Find projects in the OMOS drive (one project = one top-level folder). ALWAYS call
    this first when the user asks about any project, BRD, timeline, design, system flow,
    DB, API, or เอกสารโปรเจค.

    The drive holds hundreds of projects, so PASS A FILTER: a word from what the user
    said (Thai or English, matched anywhere in the name). Without one you only get a
    small sample, not the whole drive. Then call omos_list with the exact name you matched.

    Args:
        filter: keyword to match against project names — pass this whenever the user
            named or hinted at a project
        limit: how many names to return (default 10; raise it only when the user
            explicitly asks for the full list of projects)
    """
    import anyio.to_thread

    return await anyio.to_thread.run_sync(_omos_index, filter, limit)


MAX_INDEX_LIMIT = 500


def _omos_index(filter: str = "", limit: int = 10) -> str:
    _ensure_projects()
    all_names = sorted(_projects["by_name"])
    limit = max(1, min(limit, MAX_INDEX_LIMIT))
    needle = filter.strip().lower()
    names = [n for n in all_names if needle in n.lower()] if needle else all_names

    if needle and not names:
        return (
            f"No project name contains '{filter}' ({len(all_names)} projects total). "
            "Try a shorter or different keyword, or ask the user for the exact project name."
        )

    shown = names[:limit]
    listing = "\n".join(
        f"- {n}" + ("  ⚠️ (duplicate name)" if len(_projects["by_name"][n]) > 1 else "")
        for n in shown
    )
    if needle:
        header = f"# Projects matching '{filter}' ({len(names)} of {len(all_names)})"
        more = (
            f"\n\n_Showing {len(shown)} of {len(names)} matches — raise limit or use a "
            "more specific keyword to see the rest._" if len(names) > len(shown) else ""
        )
    else:
        header = f"# OMOS projects — sample of {len(shown)} (out of {len(all_names)})"
        more = (
            f"\n\n⚠️ _These {len(shown)} are only a sample of {len(all_names)} projects, NOT the "
            "whole drive. To find a specific project call omos_index with filter=<keyword from "
            "the user>. Only raise limit if the user really wants the full list._"
        ) if len(all_names) > len(shown) else ""

    return (
        f"{header}\n\n{listing}{more}\n\n"
        "_Then call omos_list(project) for a project's files, or omos_search(query, project) "
        "to search inside one._"
    )


@mcp.tool()
async def omos_list(project: str, subfolder: str = "") -> str:
    """List the files and folders inside ONE project, with file ids and Drive links.
    Call this after omos_index, using an exact project name from it.

    Args:
        project: exact project name from omos_index
        subfolder: optional folder name inside the project to list only that part
    """
    import anyio.to_thread

    return await anyio.to_thread.run_sync(_omos_list, project, subfolder)


MAX_LIST_ITEMS = 500


def _omos_list(project: str, subfolder: str) -> str:
    _ensure_projects()
    ids = _projects["by_name"].get(project.strip())
    if not ids:
        close = [n for n in _projects["by_name"] if project.strip().lower() in n.lower()]
        hint = ("\nDid you mean: " + ", ".join(sorted(close)[:10])) if close else ""
        return f"Unknown project '{project}'. Call omos_index for the list.{hint}"
    if len(ids) > 1:
        return (
            f"'{project}' matches {len(ids)} different folders with the same name. "
            "Ask the user which one they mean; ids: " + ", ".join(ids)
        )

    lines: list[str] = []
    truncated = ""
    deadline = time.monotonic() + DEADLINE

    def walk(folder_id: str, prefix: str, depth: int) -> None:
        nonlocal truncated
        if truncated:
            return
        if time.monotonic() > deadline:
            truncated = "time"
            return
        kids = _list_children(folder_id)
        folders = [f for f in kids if f["mimeType"] == FOLDER_MT]
        for f in kids:
            if len(lines) >= MAX_LIST_ITEMS:
                truncated = "size"
                return
            indent = "  " * depth
            if f["mimeType"] == FOLDER_MT:
                _folders[f["id"]] = (f["name"].strip(), folder_id)
                lines.append(f"{indent}- 📁 {f['name']}")
            else:
                lines.append(
                    f"{indent}- 📄 {f['name']} — id: `{f['id']}` — [link]({f.get('webViewLink', '')})"
                )
        if depth < 6:  # ponytail: depth cap, matches the drive's real nesting
            for f in folders:
                walk(f["id"], f"{prefix} / {f['name']}", depth + 1)

    root = ids[0]
    if subfolder:
        match = next(
            (f for f in _list_children(root)
             if f["mimeType"] == FOLDER_MT and f["name"].strip().lower() == subfolder.strip().lower()),
            None,
        )
        if match is None:
            return f"'{subfolder}' not found in project '{project}'. Call omos_list('{project}') to see its folders."
        root = match["id"]

    walk(root, project, 0)
    header = f"# {project}" + (f" / {subfolder}" if subfolder else "")
    if not lines:
        return f"{header}\n\n(empty)"
    if truncated == "size":
        note = (f"\n\n⚠️ _INCOMPLETE: stopped at {MAX_LIST_ITEMS} items — this is NOT the whole "
                "project. Narrow it with subfolder=, or use omos_search(query, project)._")
    elif truncated == "time":
        note = (f"\n\n⚠️ _INCOMPLETE: stopped after {DEADLINE:.0f}s — this is NOT the whole "
                "project. Narrow it with subfolder=, or use omos_search(query, project)._")
    else:
        note = ""
    return f"{header}\n\n" + "\n".join(lines) + note


@mcp.tool()
async def omos_search(query: str, project: str = "") -> str:
    """Full-text search across the OMOS drive (file content and names). Returns matching
    files with their project path, id and link; read one with omos_read.

    Args:
        query: keyword or phrase
        project: optional exact project name from omos_index to search only inside it
    """
    import anyio.to_thread

    return await anyio.to_thread.run_sync(_omos_search, query, project)


def _omos_search(query: str, project: str) -> str:
    _ensure_projects()
    project = project.strip()
    if project and project not in _projects["by_name"]:
        return f"Unknown project '{project}'. Call omos_index for the list."

    deadline = time.monotonic() + DEADLINE
    q = query.replace("\\", "\\\\").replace("'", "\\'")
    files, token = [], None
    while True:
        resp = _svc().files().list(
            q=f"fullText contains '{q}' and trashed=false and mimeType != '{FOLDER_MT}'",
            fields=f"nextPageToken, files({_FILE_FIELDS})",
            pageSize=100,
            pageToken=token,
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        files += resp.get("files", [])
        token = resp.get("nextPageToken")
        if not token or len(files) >= 200 or time.monotonic() > deadline:
            break

    results = []
    timed_out = False
    for f in files:
        if time.monotonic() > deadline:
            timed_out = True
            break
        path, proj = _path_of(f["name"], f.get("parents"))
        if project and proj != project:
            continue
        results.append(f"- 📄 {path} — id: `{f['id']}` — [link]({f.get('webViewLink', '')})")
        if len(results) >= 20:
            break

    scope = f"project '{project}'" if project else "all projects"
    if not results:
        hint = " (search was cut short by the time limit — try a narrower query or a project)" if timed_out else ""
        return f"No files matching '{query}' in {scope}{hint}. Try another keyword or omos_list."
    note = (f"\n\n⚠️ _Partial: search stopped after {DEADLINE:.0f}s — narrow the query or pass a "
            "project to see the rest._") if timed_out else ""
    return f"Found {len(results)} file(s) matching '{query}' in {scope}:\n" + "\n".join(results) + note


@mcp.tool()
async def omos_read(file_id: str):
    """Read a file from the OMOS drive as text (Google Docs/Sheets/Slides, PDF, docx, xlsx,
    markdown/text) or as an image (png/jpg). The response starts with a citation block
    (file name, path, link) — always include that citation in your answer.

    Args:
        file_id: the Drive file id from omos_index or omos_search
    """
    import anyio.to_thread

    return await anyio.to_thread.run_sync(_omos_read, file_id)


def _omos_read(file_id: str):
    try:
        return _read(file_id)
    except (TimeoutError, OSError) as exc:  # socket/HTTP timeouts from the Drive client
        link = f"https://drive.google.com/file/d/{file_id}/view"
        return (
            f"Error: reading this file timed out after {HTTP_TIMEOUT:.0f}s ({exc}). "
            f"It may be very large. Tell the user to open it directly: {link}"
        )


def _read(file_id: str):
    svc = _svc()
    meta = svc.files().get(
        fileId=file_id, fields="id,name,mimeType,webViewLink,size,parents", supportsAllDrives=True
    ).execute()
    mt, name = meta["mimeType"], meta["name"]
    cite = _cite(name, meta.get("parents"), meta.get("webViewLink", ""))

    if int(meta.get("size", 0)) > MAX_DOWNLOAD_BYTES:
        return cite + f"Error: file is too large to read ({meta['size']} bytes). Open it via the link."

    # Google-native files: export via Drive API
    exports = {
        "application/vnd.google-apps.document": "text/markdown",
        "application/vnd.google-apps.spreadsheet": "text/csv",
        "application/vnd.google-apps.presentation": "text/plain",
    }
    if mt in exports:
        try:
            data = svc.files().export(fileId=file_id, mimeType=exports[mt]).execute()
        except Exception:
            data = svc.files().export(fileId=file_id, mimeType="text/plain").execute()
        return cite + _truncate(data.decode("utf-8", errors="replace"))
    if mt.startswith("application/vnd.google-apps"):
        return cite + f"Error: unsupported Google file type '{mt}'. Open it via the link."

    data = svc.files().get_media(fileId=file_id, supportsAllDrives=True).execute()

    if mt.startswith("image/"):
        fmt = mt.split("/", 1)[1].split("+")[0]
        return [cite + "(image below)", Image(data=data, format=fmt)]
    if mt == "application/pdf":
        from pypdf import PdfReader

        pages = [page.extract_text() or "" for page in PdfReader(io.BytesIO(data)).pages]
        return cite + _truncate("\n\n".join(pages))
    if name.lower().endswith(".docx"):
        return cite + _truncate(_read_docx(data))
    if name.lower().endswith((".xlsx", ".xlsm")):
        return cite + _truncate(_read_xlsx(data))
    if mt.startswith("text/") or mt in ("application/json", "application/xml") or name.lower().endswith(
        (".md", ".txt", ".csv", ".json", ".yaml", ".yml")
    ):
        return cite + _truncate(data.decode("utf-8", errors="replace"))

    return cite + f"Error: unsupported file type '{mt}'. Open it via the link."


@mcp.tool()
async def omos_refresh() -> str:
    """Reload the project list so newly added projects show up immediately
    (it is otherwise cached for a few minutes)."""
    import anyio.to_thread

    _projects["built_at"] = 0.0
    _folders.clear()
    await anyio.to_thread.run_sync(_ensure_projects)
    return f"Project list reloaded: {len(_projects['by_name'])} project(s)."


# --- File converters ---


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS] + f"\n\n[... truncated, {len(text)} chars total]"
    return text or "(file is empty)"


def _read_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        parts.append(
            "\n".join(" | ".join(c.text.strip() for c in row.cells) for row in table.rows)
        )
    return "\n\n".join(parts)


def _read_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 200:  # ponytail: row cap per sheet, raise if real sheets are bigger
                rows.append("[... more rows truncated]")
                break
            rows.append(" | ".join("" if v is None else str(v) for v in row))
        parts.append(f"### Sheet: {ws.title}\n" + "\n".join(rows))
    return "\n\n".join(parts)


def main():
    mcp.run(transport="stdio")


# --- Remote HTTP mode (same skeleton as wiki-mcp) ---


class _OAuthTokenVerifier:
    """Validates bearer tokens issued by an external OAuth provider (WorkOS AuthKit, Auth0, …).

    Accepts either a valid provider JWT (so Claude.ai / ChatGPT web work via OAuth)
    or the shared team token (so header-capable clients keep working).
    """

    def __init__(self, issuer: str, audience: str, static_token: str = ""):
        import jwt  # provided transitively by mcp

        self._jwt = jwt
        self._issuer = issuer
        # Accept the audience with and without a trailing slash: providers and the
        # MCP resource-metadata normalization don't always agree on it.
        base = audience.rstrip("/")
        self._audience = [base, base + "/"]
        self._static = static_token
        self._jwk_client = jwt.PyJWKClient(self._discover_jwks(issuer))

    @staticmethod
    def _discover_jwks(issuer: str) -> str:
        import urllib.request

        base = issuer.rstrip("/")
        for path in ("/.well-known/openid-configuration", "/.well-known/oauth-authorization-server"):
            try:
                with urllib.request.urlopen(base + path, timeout=10) as resp:  # noqa: S310
                    return json.load(resp)["jwks_uri"]
            except Exception:
                continue
        raise RuntimeError(f"Could not discover JWKS endpoint from issuer {issuer}")

    async def verify_token(self, token: str):
        import hmac
        import logging

        from mcp.server.auth.provider import AccessToken

        log = logging.getLogger("omos_mcp.auth")

        if self._static and hmac.compare_digest(token, self._static):
            return AccessToken(token=token, client_id="team-token", scopes=[], subject="team")
        try:
            key = self._jwk_client.get_signing_key_from_jwt(token).key
            # Verify signature + audience here; check issuer manually below so a
            # trailing-slash difference doesn't cause a spurious rejection.
            claims = self._jwt.decode(
                token, key,
                algorithms=["RS256", "RS384", "RS512"],
                audience=self._audience,
                options={"verify_iss": False},
            )
        except Exception as exc:
            log.warning("omos-mcp: token rejected during decode: %s", exc)
            return None

        token_iss = (claims.get("iss") or "").rstrip("/")
        if token_iss != self._issuer.rstrip("/"):
            log.warning("omos-mcp: issuer mismatch: token=%r expected=%r", token_iss, self._issuer)
            return None

        scope = claims.get("scope", "")
        return AccessToken(
            token=token,
            client_id=claims.get("azp") or claims.get("client_id") or "oauth",
            scopes=scope.split() if isinstance(scope, str) else list(scope or []),
            subject=claims.get("sub"),
            expires_at=claims.get("exp"),
            resource=self._audience[0],
            claims=claims,
        )


def _build_http_server(**kwargs) -> FastMCP:
    """Build a FastMCP for HTTP serving with the OMOS tools registered.

    DNS-rebinding protection is disabled because the server runs behind a TLS
    reverse proxy (e.g. Render) under an external hostname and is already
    protected by auth — the default localhost-only Host allowlist would
    otherwise reject every proxied request with 421.
    """
    from mcp.server.transport_security import TransportSecuritySettings

    server = FastMCP(
        "omos-mcp",
        instructions=INSTRUCTIONS,
        transport_security=TransportSecuritySettings(enable_dns_rebinding_protection=False),
        **kwargs,
    )
    for fn in (omos_index, omos_list, omos_search, omos_read, omos_refresh):
        server.tool()(fn)
    return server


def main_http():
    """Run as a remote HTTP (streamable) server.

    Two auth modes (auto-selected by env):
      * OAuth — set OAUTH_ISSUER, OAUTH_AUDIENCE, PUBLIC_URL. Web clients
        (Claude.ai, ChatGPT) authenticate through the provider. The shared
        OMOS_AUTH_TOKEN, if set, still works for header-capable clients.
      * Shared bearer — set only OMOS_AUTH_TOKEN. Works with header-capable
        clients (Claude Code, Cursor, VS Code), not web OAuth clients.
    """
    import hmac

    import uvicorn
    from starlette.middleware.base import BaseHTTPMiddleware
    from starlette.requests import Request
    from starlette.responses import JSONResponse, PlainTextResponse
    from starlette.routing import Route

    token = os.environ.get("OMOS_AUTH_TOKEN", "")
    issuer = os.environ.get("OAUTH_ISSUER", "")
    audience = os.environ.get("OAUTH_AUDIENCE", "")
    public_url = os.environ.get("PUBLIC_URL", "")

    _svc()  # build Drive client at startup (fails loudly on bad credentials)

    async def health(_request: Request) -> PlainTextResponse:
        return PlainTextResponse("ok")

    if issuer:
        if not (audience and public_url):
            raise SystemExit(
                "OAUTH_ISSUER is set, so OAUTH_AUDIENCE and PUBLIC_URL are required too."
            )
        from mcp.server.auth.settings import AuthSettings

        server = _build_http_server(
            token_verifier=_OAuthTokenVerifier(issuer, audience, token),
            auth=AuthSettings(
                issuer_url=issuer,
                resource_server_url=public_url,
                required_scopes=[],
            ),
        )
        app = server.streamable_http_app()
        app.router.routes.append(Route("/healthz", health, methods=["GET"]))
    else:
        if not token:
            raise SystemExit(
                "Set OMOS_AUTH_TOKEN (shared-bearer mode) or OAUTH_ISSUER + "
                "OAUTH_AUDIENCE + PUBLIC_URL (OAuth mode) so the endpoint is not public."
            )

        class BearerAuth(BaseHTTPMiddleware):
            async def dispatch(self, request: Request, call_next):
                if request.url.path == "/healthz":
                    return await call_next(request)
                header = request.headers.get("authorization", "")
                provided = header[7:] if header.startswith("Bearer ") else ""
                if not hmac.compare_digest(provided, token):
                    return JSONResponse({"error": "unauthorized"}, status_code=401)
                return await call_next(request)

        app = _build_http_server().streamable_http_app()
        app.router.routes.append(Route("/healthz", health, methods=["GET"]))
        app.add_middleware(BearerAuth)

    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", "8000")))


if __name__ == "__main__":
    main()
